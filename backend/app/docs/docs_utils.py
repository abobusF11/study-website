from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Cm
from typing import List, Dict, Any

def merge_vertical_cells(table, start_row, end_row, col):
    """Объединяет ячейки вертикально"""
    for row in range(start_row, end_row):
        table.cell(row, col).merge(table.cell(row+1, col))

def replace_text_in_docx(doc, replacements):
    """Заменяет текст в документе"""
    for paragraph in doc.paragraphs:
        for old_text, new_text in replacements.items():
            if old_text in paragraph.text:
                paragraph.text = paragraph.text.replace(old_text, new_text)

def add_rows_to_table(doc, input_data: List[Dict[str, Any]], cell_config: List[Dict[str, Any]]):
    """Добавляет строки в таблицу документа на основе конфигурации"""
    if not doc.tables:
        return  # Если таблиц нет, просто выходим из функции

    table = doc.tables[0]

    # Удаляем все строки с данными, оставляя только заголовок
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[1]._tr)

    # Фиксированные параметры
    DEFAULT_ROW_HEIGHT = Pt(24)
    FONT_SIZE = Pt(11)
    COLUMN_WIDTH = Cm(2.5)

    # Группируем данные по организации
    org_groups = {}
    for idx, data in enumerate(input_data):
        org = data.get("org", "")  # Исправлено на прямое обращение к полю
        org_groups.setdefault(org, []).append((idx, data))

    # Добавляем строки
    for org, group_data in org_groups.items():
        group_rows = []

        # Рассчитываем высоту строки
        row_height = DEFAULT_ROW_HEIGHT
        max_text_length = max(len(str(data.get(key, ""))) for _, data in group_data for key in data)
        if max_text_length > 50:
            row_height = Pt(36)

        for i, (data_idx, data) in enumerate(group_data):
            row = table.add_row()
            row.height = row_height
            group_rows.append(row)

            for col_idx, config in enumerate(cell_config):
                if "type" not in config:
                    continue  # Пропускаем некорректные конфиги

                cell = row.cells[col_idx]
                cell.text = ""

                try:
                    if config["type"] == "index":
                        cell.text = str(data_idx + 1)
                    elif config["type"] == "data":
                        if "key" not in config:
                            continue

                        # Используем ключ напрямую, так как в prepare_dict_data уже произошло преобразование
                        field_key = config["key"]
                        cell.text = str(data.get(field_key, ""))
                    elif config["type"] == "static":
                        cell.text = str(config.get("value", ""))
                except Exception as e:
                    continue

                # Форматирование ячейки
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.0
                    for run in paragraph.runs:
                        run.font.size = FONT_SIZE
                        run.font.italic = (config.get("style") == "italic")

        # Объединение ячеек
        if len(group_rows) > 1:
            for col_idx, config in enumerate(cell_config):
                if config.get("type") == "data" and config.get("key") == "Наименование предприятия":
                    group_rows[0].cells[col_idx].merge(group_rows[-1].cells[col_idx])

    # Устанавливаем ширину столбцов
    for col in table.columns:
        col.width = COLUMN_WIDTH

FIELD_NAME_MAPPING = {
    "initials": "ФИО",
    "inn": "ИНН сотрудника",
    "org": "Наименование предприятия",
    "safety": "Группа по безопасности",
    "reg_num": "Регистрационный номер",
    "position": "Занимаемая должность",
    "org_inn": "ИНН Организации",
    "snils": "СНИЛС"
}

def get_field_display_name(field: str) -> str:
    """Возвращает отображаемое имя поля"""
    return FIELD_NAME_MAPPING.get(field, field)