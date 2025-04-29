from datetime import date, datetime
from pathlib import Path
from typing import List
from copy import deepcopy
from docx import *
from fastapi import Depends, Query, HTTPException, APIRouter
from sqlalchemy import *
from sqlalchemy.exc import IntegrityError
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import joinedload
from starlette.responses import FileResponse

from backend.database import get_db
from backend.models import TeachersGroup
from backend.models.group_model import Group, Client, CourseGroup
from .schemas import GroupCreate, ErrorResponse, GroupUpdate, GroupCreateResponse, GroupResponse, CourseResponse, \
    ClientResponse, GroupUpdateResponse
from .utils import _process_courses, COURSE_HOURS, calculate_end_date, _process_teachers
from ..teachers.schemas import TeacherResponse
from ...auth.utils import require_lvl

router = APIRouter(tags=["Groups"])


@router.post("/create",
             response_model=GroupCreateResponse,
             responses={400: {"model": ErrorResponse}, 500: {"model": ErrorResponse}})
async def create_group(
        group_data: GroupCreate,
        db: AsyncSession = Depends(get_db),
        _=Depends(require_lvl(1))
):
    # 1. Создаем группу
    group = Group(date=group_data.date)
    db.add(group)
    await db.flush()  # Получаем ID группы

    if (group_data.teachers):
        for i in group_data.teachers:
            staff = TeachersGroup(
                teacher_id=i,
                group_id=group.id
            )
            db.add(staff)

    # 2. Создаем курсы для этой группы
    for course in group_data.courses:
        course_group = CourseGroup(
            course_id=course.course_id,
            group_id=group.id  # Используем ID только что созданной группы
        )
        db.add(course_group)
        await db.flush()  # Получаем ID course_group

        # 3. Создаем клиентов для этого курса
        for client in course.clients:
            client_record = Client(
                initials=client.initials,
                inn=client.inn,
                org=client.org,
                safety=client.safety,
                reg_num=client.reg_num,
                course_group_id=course_group.id
            )
            db.add(client_record)

    await db.commit()

    return {
        "id": group.id
    }


@router.get(
    "/show",
    response_model=List[GroupResponse],
    responses={401: {"model": ErrorResponse}, 404: {"model": ErrorResponse}}
)
async def show_groups(
        db: AsyncSession = Depends(get_db),
        date_filter: str = Query(
            None,
            description="Фильтр по дате: 'active', 'archive', 'from-user'",
            regex="^(active|archive|from-user)$"
        )
        # _=Depends(require_lvl(2))
):
    # Загружаем группы с их курсами и клиентами
    result = await db.execute(
        select(Group)
        .options(
            joinedload(Group.courses)
            .joinedload(CourseGroup.clients),
            joinedload(Group.teacher_groups)
            .joinedload(TeachersGroup.teacher)
        )
        .order_by(Group.date.desc())  # Сортировка по дате (новые сначала)
    )

    groups = result.unique().scalars().all()

    if not groups:
        raise HTTPException(status_code=404, detail="Нет групп")

    current_date = date.today()
    filtered_groups = []

    for group in groups:
        # Находим курс с максимальным количеством часов в группе
        max_hours_course = max(
            group.courses,
            key=lambda c: COURSE_HOURS.get(c.course_id, 0),
            default=None
        )

        if max_hours_course:
            # Рассчитываем дату окончания для курса с максимальными часами
            end_date = calculate_end_date(group.date, max_hours_course.course_id)  # Преобразуем в date
            # Применяем фильтр
            if not date_filter:
                filtered_groups.append(group)
            elif date_filter == "active" and end_date >= current_date:
                filtered_groups.append(group)
            elif date_filter == "archive" and end_date < current_date:
                filtered_groups.append(group)

    if not filtered_groups:
        raise HTTPException(status_code=404, detail="Нет активных групп")

    return [
        GroupResponse(
            id=group.id,
            date=group.date,
            courses=[
                CourseResponse(
                    id=course.id,
                    course_id=course.course_id,
                    group_id=course.group_id,
                    clients=sorted([
                        ClientResponse(
                            id=client.id,
                            initials=client.initials,
                            inn=client.inn,
                            org=client.org,
                            safety=client.safety,
                            reg_num=client.reg_num
                        )
                        for client in course.clients
                    ], key=lambda x: x.initials)
                )
                for course in group.courses
            ],
            teachers=[
                TeacherResponse(
                    id=teacher_group.teacher.id,
                    initials=teacher_group.teacher.initials,
                    status=teacher_group.teacher.status
                )
                for teacher_group in group.teacher_groups
            ]
        )
        for group in filtered_groups
    ]


@router.put(
    "/update",
    response_model=GroupUpdateResponse,
    responses={
        400: {"model": ErrorResponse},
        401: {"model": ErrorResponse},
        404: {"model": ErrorResponse},
        422: {"model": ErrorResponse},
        500: {"model": ErrorResponse}
    }
)
async def update_group(
        group_update: GroupUpdate,
        db: AsyncSession = Depends(get_db),
        _=Depends(require_lvl(1))
):
    try:
        async with db.begin():
            # Получаем группу со всеми связанными данными
            group = await db.execute(
                select(Group)
                .where(Group.id == group_update.id)
                .options(
                    joinedload(Group.courses)
                    .joinedload(CourseGroup.clients),
                    joinedload(Group.teacher_groups)
                )
            )
            group = group.scalars().first()

            if not group:
                raise HTTPException(status_code=404, detail="Group not found")

            # Обновляем основные поля группы
            group.date = group_update.date

            # Обрабатываем курсы группы
            await _process_courses(db, group, group_update.courses)

            await _process_teachers(db, group, group_update.teachers)

            return GroupUpdateResponse(id=group.id)

    except IntegrityError as e:
        if "unique constraint" in str(e).lower():
            raise HTTPException(
                status_code=400,
                detail="Duplicate data violation"
            )
        raise HTTPException(
            status_code=400,
            detail="Database integrity error"
        )
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Internal server error {e}"
        )


@router.delete(
    "/delete",
    responses={401: {"model": ErrorResponse}, 404: {"model": ErrorResponse}}
)
async def delete_group(
        group_id: int = Query(..., description="ID группы для удаления"),
        db: AsyncSession = Depends(get_db),
        _=Depends(require_lvl(1))
):
    result = await db.execute(
        select(Group)
        .where(Group.id == group_id)
    )
    group = result.scalar_one_or_none()

    if not group:
        raise HTTPException(status_code=404, detail="Group not found")

    await db.delete(group)
    await db.commit()

    return {"message": f"Group {group_id} and all its clients were deleted"}


TEMPLATES_DIR = Path(__file__).parent.parent.parent / "template" / "group" / "docs"
TEMP_DIR = Path(__file__).parent.parent.parent / "temp"

def replace_text(paragraph, old_text, new_text):
    if old_text in paragraph.text:
        inline = paragraph.runs
        for i in range(len(inline)):
            if old_text in inline[i].text:
                inline[i].text = inline[i].text.replace(old_text, new_text)

# @router.get("/{id_group}")
# async def get_template(id_group: int, db: AsyncSession = Depends(get_db)):
#     # Получаем группу и связанных клиентов
#     result = await db.execute(
#         select(Group).where(Group.id == id_group)
#         .options(
#             joinedload(Group.courses)
#             .joinedload(CourseGroup.clients)
#         )
#     )
#     group = result.unique().scalars().first()
#
#     if not group:
#         raise HTTPException(status_code=404, detail="Группа не найдена")
#
#     if not group.courses:
#         raise HTTPException(status_code=400, detail="В группе нет курсов")
#
#     # Берем первый курс (можно доработать для нескольких курсов)
#     course = group.courses[0]
#
#     # Определяем шаблон в зависимости от course_id
#     template_mapping = {
#         1: "1.docx",  # Для курса по работе на высоте
#         2: "1.docx", # Для курса по электробезопасности
#         3: "1.docx", # Для курса по электробезопасности
#         # Добавьте другие курсы по необходимости
#     }
#
#     template_name = template_mapping.get(course.course_id, "template_default.docx")
#     template_path = TEMPLATES_DIR / template_name
#
#     if not template_path.exists():
#         raise HTTPException(status_code=404, detail=f"Шаблон {template_name} не найден")
#
#     try:
#         doc = Document(template_path)
#
#         # Заменяем общие данные
#         for paragraph in doc.paragraphs:
#             replace_text(paragraph, "{{date}}", group.date.strftime("%d.%m.%Y"))
#             replace_text(paragraph, "{{protocol_num}}", str(group.id))
#
#         # Обрабатываем таблицу
#         for table in doc.tables:
#             if len(table.rows) < 3:
#                 continue
#
#             sample_row = table.rows[1]
#             for idx, client in enumerate(course.clients, start=1):
#                 new_row = deepcopy(sample_row)
#
#                 for cell in new_row.cells:
#                     for paragraph in cell.paragraphs:
#                         replace_text(paragraph, "{{client_name}}", client.initials)
#                         replace_text(paragraph, "{{org}}", client.org)
#                         replace_text(paragraph, "{{position}}", getattr(client, "position", ""))
#                         replace_text(paragraph, "{{safety_group}}", str(client.safety) if client.safety else "")
#
#                 table._tbl.append(new_row._tr)
#
#             # Удаляем строку-образец после обработки
#             table._tbl.remove(sample_row._tr)
#
#         # Сохраняем результат
#         TEMP_DIR.mkdir(parents=True, exist_ok=True)
#         temp_filename = f"protocol_{group.id}_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
#         temp_path = TEMP_DIR / temp_filename
#         doc.save(temp_path)
#
#         return FileResponse(
#             temp_path,
#             filename=temp_filename,
#             media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
#         )
#
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=str(e))